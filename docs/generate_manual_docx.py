from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


WORKSPACE = Path("/Users/user/eclipse-workspace/badminton-court-management")
MD_PATH = WORKSPACE / "docs" / "Huong_dan_su_dung_BadmintonCourtManagement.md"
DOCX_PATH = WORKSPACE / "Huong_dan_su_dung_BadmintonCourtManagement.docx"

IMAGE_PATTERN = re.compile(r"!\[(.*?)\]\((.*?)\)")


def set_font(style, name="Times New Roman", size=None):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        style.font.size = size


def build_docx():
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()

    doc = Document()
    styles = doc.styles
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        set_font(styles[style_name])
    if "Caption" in styles:
        set_font(styles["Caption"])
    styles["Normal"].font.size = Pt(12)

    number_buffer = []

    def flush_number_buffer():
        nonlocal number_buffer
        if not number_buffer:
            return
        for item in number_buffer:
            p = doc.add_paragraph(style="List Number")
            p.add_run(item)
        number_buffer = []

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_number_buffer()
            doc.add_paragraph("")
            continue

        image_match = IMAGE_PATTERN.fullmatch(stripped)
        if image_match:
            flush_number_buffer()
            caption, rel_path = image_match.groups()
            image_path = (MD_PATH.parent / rel_path).resolve()
            if image_path.exists():
                pic = doc.add_paragraph()
                pic.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                pic.add_run().add_picture(str(image_path), width=Inches(6.8))

                cap = doc.add_paragraph(caption)
                cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if cap.runs:
                    cap.runs[0].italic = True
            continue

        if line.startswith("# "):
            flush_number_buffer()
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.add_run(line[2:].strip())
            continue

        if line.startswith("## "):
            flush_number_buffer()
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
            continue

        if line.startswith("### "):
            flush_number_buffer()
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            continue

        if line.startswith("- "):
            flush_number_buffer()
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:].strip())
            continue

        parts = stripped.split(". ", 1)
        if len(parts) == 2 and parts[0].isdigit():
            number_buffer.append(parts[1])
            continue

        flush_number_buffer()
        para = doc.add_paragraph(line)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    flush_number_buffer()

    for p in doc.paragraphs:
        if p.style.name != "Title":
            p.paragraph_format.space_after = Pt(6)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    print(build_docx())
